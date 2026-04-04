import os
import re
import shutil
import tempfile
import subprocess
from pathlib import Path
from typing import List

from docxtpl import DocxTemplate
from docx import Document
from copy import deepcopy

from app.models.schemas import PreferenceListItem

BASE_DIR = Path(__file__).parent.parent / "templates"

FRONT_TEMPLATE = BASE_DIR / "F.docx"
TABLE_TEMPLATE = BASE_DIR / "Table.docx"


# ─────────────────────────────────────────────
# ✅ STEP 1: GENERATE FRONT PAGE (JINJA)
# ─────────────────────────────────────────────

def generate_front_page(
    student_name,
    counselling_id,
    percentile,
    category,
    rank,
    mobile,
    preferred_cities,
    preferred_branches,
):
    doc = DocxTemplate(str(FRONT_TEMPLATE))

    cities = {
        f"city_{i+1}": preferred_cities[i] if i < len(preferred_cities) else ""
        for i in range(6)
    }

    branches = {
        f"branch_{i+1}": preferred_branches[i] if i < len(preferred_branches) else ""
        for i in range(6)
    }

    context = {
        "Student_name": student_name,
        "Id": counselling_id,
        "percentile": f"{percentile:.4f}",
        "category": category,
        "rank": rank,
        "mobile": mobile,
        **cities,
        **branches,
    }

    doc.render(context)

    tmp_dir = tempfile.mkdtemp()
    front_path = os.path.join(tmp_dir, "front.docx")
    doc.save(front_path)

    return front_path, tmp_dir


# ─────────────────────────────────────────────
# ✅ STEP 2: GENERATE TABLE PAGE
# ─────────────────────────────────────────────

def generate_table_page(preference_list: List[PreferenceListItem], tmp_dir):

    doc = Document(str(TABLE_TEMPLATE))
    table = doc.tables[0]

    # remove all rows except header
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for idx, item in enumerate(preference_list):
        row = table.add_row().cells

        row[0].text = str(idx + 1)
        row[1].text = item.college_code
        row[2].text = item.college_name
        row[3].text = item.city or ""
        row[4].text = item.category
        row[5].text = item.branch_code
        row[6].text = item.branch_name
        row[7].text = f"{item.cutoff_percentile:.4f}"

    table_path = os.path.join(tmp_dir, "table.docx")
    doc.save(table_path)

    return table_path


# ─────────────────────────────────────────────
# ✅ STEP 3: MERGE DOCX FILES
# ─────────────────────────────────────────────

def merge_docs(front_path, table_path, tmp_dir):

    main_doc = Document(front_path)
    table_doc = Document(table_path)

    # Add page break
    main_doc.add_page_break()

    # Append all elements from table doc
    for element in table_doc.element.body:
        main_doc.element.body.append(deepcopy(element))

    merged_path = os.path.join(tmp_dir, "final.docx")
    main_doc.save(merged_path)

    return merged_path


# ─────────────────────────────────────────────
# ✅ STEP 4: CONVERT TO PDF
# ─────────────────────────────────────────────

def convert_to_pdf(docx_path, tmp_dir):

    soffice = shutil.which("soffice")

    if not soffice:
        soffice = r"C:\Program Files\LibreOffice\program\soffice.exe"

    if not os.path.exists(soffice):
        raise RuntimeError("LibreOffice not found")

    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", tmp_dir,
            docx_path,
        ],
        check=True,
    )

    pdf_path = docx_path.replace(".docx", ".pdf")

    if not os.path.exists(pdf_path):
        raise RuntimeError("PDF not generated")

    return pdf_path


# ─────────────────────────────────────────────
# ✅ FINAL MAIN FUNCTION
# ─────────────────────────────────────────────

def fill_template(
    student_name: str,
    counselling_id: str,
    percentile: float,
    category: str,
    rank: str,
    mobile: str,
    preferred_cities: List[str],
    preferred_branches: List[str],
    preference_list: List[PreferenceListItem],
) -> str:

    # 1️⃣ Front page
    front_path, tmp_dir = generate_front_page(
        student_name,
        counselling_id,
        percentile,
        category,
        rank,
        mobile,
        preferred_cities,
        preferred_branches,
    )

    # 2️⃣ Table page
    table_path = generate_table_page(preference_list, tmp_dir)

    # 3️⃣ Merge
    merged_docx = merge_docs(front_path, table_path, tmp_dir)

    # 4️⃣ Convert
    pdf_path = convert_to_pdf(merged_docx, tmp_dir)

    return pdf_path